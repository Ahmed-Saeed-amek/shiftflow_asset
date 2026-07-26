# -*- coding: utf-8 -*-
"""Part 1: Title page, TOC, Introduction, Overall Description, Architecture.
Imported and called by build_srs.py's main assembly step.
"""


def build_part1(ctx):
    doc = ctx["doc"]
    h1_ = ctx["h1"]; h2 = ctx["h2"]; h3 = ctx["h3"]
    para = ctx["para"]; bullet = ctx["bullet"]; numbered = ctx["numbered"]
    add_diagram = ctx["add_diagram"]; simple_table = ctx["simple_table"]
    page_break = ctx["page_break"]; add_toc = ctx["add_toc"]
    NAVY = ctx["NAVY"]; ACCENT = ctx["ACCENT"]; GRAY = ctx["GRAY"]
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # ---------------- Title page ----------------
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Software Requirements Specification")
    r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = NAVY

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("ShiftFlow")
    r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = ACCENT

    s2 = doc.add_paragraph()
    s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s2.add_run("Shift Scheduling, Workforce Operations & Asset/Vendor Maintenance Management System")
    r.font.size = Pt(13); r.font.italic = True; r.font.color.rgb = GRAY

    for _ in range(6):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Version 1.0")
    r.font.size = Pt(12); r.font.color.rgb = GRAY

    page_break()

    # ---------------- Document control ----------------
    h2("Document Control")
    simple_table(
        ["Field", "Value"],
        [
            ["Document Title", "Software Requirements Specification — ShiftFlow"],
            ["Version", "1.0"],
            ["Status", "Draft"],
            ["Classification", "Internal"],
            ["Prepared For", "ShiftFlow Project Stakeholders"],
            ["Scope of This Document", "Covers the ShiftFlow.Web application and its ShiftFlow.Mobile companion shell, as implemented at the time of writing."],
        ],
        widths=[1.8, 4.5],
    )

    h3("Revision History")
    simple_table(
        ["Version", "Description"],
        [["1.0", "Initial comprehensive SRS covering all implemented modules: Identity/RBAC, Shift Scheduling, Shift Operations, Change Requests, Reports & Analytics, AI Assistant, Asset Management, Work Orders, Vendor Portal, Contracts, Localization, and the Mobile companion app."]],
        widths=[1.2, 5.1],
    )

    page_break()

    # ---------------- TOC ----------------
    h1_("Table of Contents")
    add_toc(doc)
    page_break()

    # ================= 1. INTRODUCTION =================
    h1_("1. Introduction")

    h2("1.1 Purpose")
    para(
        "This Software Requirements Specification (SRS) describes the functional and non-functional "
        "requirements of ShiftFlow, a bilingual (Arabic/English) web application that manages shift "
        "scheduling, live shift operations, workforce change requests, reporting/analytics, an AI-assisted "
        "chat interface, and a full asset/vendor/work-order maintenance-management module. The document is "
        "intended for developers, testers, project stakeholders, and future maintainers who need an accurate, "
        "detailed reference for what the system does and how its major workflows behave."
    )

    h2("1.2 Scope")
    para(
        "ShiftFlow is delivered as an ASP.NET Core 8 MVC web application (“ShiftFlow.Web”), packaged "
        "additionally as a thin .NET MAUI WebView shell (“ShiftFlow.Mobile”) for native app-store "
        "distribution on Android, iOS, Mac Catalyst, and Windows. The system serves an organization that "
        "operates shift-based field/control-room work (the seeded reference data models a Kuwaiti electricity "
        "and water utility with stations in Kuwait City, Ahmadi, and Jahra), and covers:"
    )
    bullet("Shift rotation planning and schedule publication")
    bullet("Live, per-shift operational execution: activation, attendance, task handover, incident logging, and closure reporting")
    bullet("Employee-initiated and manager-initiated change requests (temporary transfers, overtime, absence, swaps, replacements, permanent group changes)")
    bullet("Executive and task-completion analytics dashboards, plus self-service personal metrics")
    bullet("A tool-calling AI chat assistant that can answer questions and perform permitted actions on the user's behalf")
    bullet("A hierarchical asset register (Governorate → Area → Zone → Asset) with categorized failure/action reporting")
    bullet("A vendor-driven work order repair workflow, including a dedicated Vendor Portal with its own login accounts")
    bullet("Vendor and contract management, including contract-to-asset linking")
    bullet("Fine-grained, cacheable role- and user-level permission (RBAC) administration")
    bullet("Full Arabic/English localization with right-to-left (RTL) layout mirroring")

    h2("1.3 Definitions, Acronyms, and Abbreviations")
    simple_table(
        ["Term", "Definition"],
        [
            ["RBAC", "Role-Based Access Control — ShiftFlow's fine-grained permission model, layered on top of ASP.NET Core Identity roles."],
            ["Work Area", "A physical/organizational site (e.g. a station) that groups shift crews."],
            ["Shift Group", "One of five fixed crew rotation groups per Work Area, labeled A, B, C, D, F."],
            ["DailyGroupShift", "The system's per-day, per-group shift-plan record — the operational source of truth once a schedule exists."],
            ["Rotation Template", "A reusable 5-day pattern defining which group works Morning/Evening/Night/Off each day."],
            ["Handover", "The act of carrying an unfinished, mandatory task forward from a closing shift onto the same group's next shift."],
            ["Change Request", "An employee- or manager-initiated request to alter a scheduled shift assignment (transfer, overtime, absence, swap, replacement, or permanent group change)."],
            ["Work Order", "A maintenance record tracking an asset repair from report through vendor resolution to closure."],
            ["Service Contract", "A Contract with ContractType = “Service” — the only contract type used to automatically resolve which vendor a work order is sent to."],
            ["Vendor Portal", "A restricted area of the application, reachable only by accounts with the “Vendor” Identity role, showing only that vendor's own work orders."],
            ["RTL / LTR", "Right-to-left / left-to-right text and layout direction, switched based on the active language (Arabic = RTL, English = LTR)."],
            ["Entra ID", "Microsoft Entra ID (formerly Azure Active Directory) — the optional external identity provider ShiftFlow can federate logins with."],
            ["SRS", "Software Requirements Specification (this document)."],
        ],
        widths=[1.6, 4.7],
    )

    h2("1.4 References")
    bullet("ShiftFlow.Web source code and inline documentation (C:\\...\\solution3\\ShiftFlow.Web)")
    bullet("ShiftFlow.Mobile source code and README (C:\\...\\solution3\\ShiftFlow.Mobile)")
    bullet("ASP.NET Core 8 / Entity Framework Core 8 official documentation")
    bullet("Microsoft Entra ID (OpenID Connect) documentation")

    h2("1.5 Document Overview")
    para(
        "Section 2 gives an overall product description. Section 3 covers system architecture and technology. "
        "Section 4 covers identity and access control. Sections 5–13 specify functional requirements module "
        "by module, each including the relevant workflow diagrams. Section 14 covers localization, Section 15 "
        "the mobile app, Section 16 the data model, Section 17 external interfaces, and Section 18 non-functional "
        "requirements. Appendices provide the full permission catalog, role-permission matrix, status/enum "
        "reference tables, and a glossary."
    )

    page_break()

    # ================= 2. OVERALL DESCRIPTION =================
    h1_("2. Overall Description")

    h2("2.1 Product Perspective")
    para(
        "ShiftFlow is a self-contained, server-rendered web application (Razor Views over ASP.NET Core MVC "
        "controllers) backed by a single SQL Server database. It is not part of a larger product family, but "
        "it does integrate with external services: Microsoft Entra ID for federated login, SMTP for email "
        "notifications, Twilio for WhatsApp notifications, and an OpenAI-compatible LLM (direct OpenAI or "
        "Azure OpenAI) plus Azure Cognitive Services Speech for its AI Assistant's voice features. The mobile "
        "app is not an independent client; it is a native WebView shell over the same web application."
    )
    add_diagram("architecture.png", "Figure 2.1 — High-level system architecture")

    h2("2.2 Product Functions (Summary)")
    bullet("Rotation-template-driven shift schedule generation, publication, manual override, and archival")
    bullet("Live shift execution: activation, attendance tracking, task management with automatic handover, incident logging, and closure reporting with file attachments")
    bullet("Change-request workflow with employee accept/decline gating for manager-initiated transfers and overtime")
    bullet("Executive dashboard, task-completion analytics, and self-service personal metrics")
    bullet("AI chat assistant with 27 permission-gated tool functions and voice/avatar support")
    bullet("Hierarchical asset register with categorized reporting and a vendor-driven, multi-stage work order repair workflow")
    bullet("Dedicated Vendor Portal with ownership-enforced access to only a vendor's own work orders")
    bullet("Vendor and contract management with category-based bulk asset linking")
    bullet("Fine-grained RBAC administration (role and per-user permission overrides)")
    bullet("Full Arabic/English bilingual UI with RTL/LTR mirroring")

    h2("2.3 User Classes and Characteristics")
    simple_table(
        ["Role", "Typical Characteristics"],
        [
            ["Admin", "Full system access; manages users, RBAC, all scheduling and asset-management functions."],
            ["ShiftManager", "Manages scheduling, shift operations, change-request review, assets, vendors, contracts, and work orders across the organization."],
            ["Supervisor", "Work-area-scoped operational management: attendance, tasks, incidents, change-request review, asset/work-order management."],
            ["Section Head", "Read-oriented oversight role: schedule and change-request visibility, asset/vendor/work-order viewing."],
            ["Senior Engineer / Engineer / Operation Engineer / Technician", "Field/operational staff: view their own schedule and tasks, update task status, submit change requests, report asset failures, and (Engineer/Senior Engineer/Operation Engineer/Technician) manage work orders."],
            ["HR", "Schedule visibility and asset visibility only — no operational or administrative actions."],
            ["Vendor", "External maintenance provider; restricted entirely to the Vendor Portal, seeing only work orders assigned to their own vendor account."],
        ],
        widths=[1.9, 4.4],
    )

    h2("2.4 Operating Environment")
    bullet("Server: ASP.NET Core 8 runtime on Windows or Linux; SQL Server (LocalDB for development, full SQL Server in production)")
    bullet("Client: any modern evergreen web browser (desktop or mobile), or the ShiftFlow.Mobile native app shell on Android, iOS, Mac Catalyst, or Windows")
    bullet("Network: HTTPS required in production (HSTS enforced outside Development); the mobile app uses plain HTTP only for local development against emulator loopback addresses")

    h2("2.5 Design and Implementation Constraints")
    bullet("Single relational database (SQL Server) via EF Core 8 code-first migrations — no polyglot persistence")
    bullet("Server-rendered Razor Views — no separate SPA/JS framework front end; JavaScript is used for targeted AJAX interactions (typeahead pickers, cascading dropdowns, chip-based multi-select)")
    bullet("File uploads (shift report and work order attachments) are stored under App_Data/uploads/*, deliberately outside wwwroot, so they are unreachable via the static-file pipeline and must be served through an authenticated, ownership-checked download action")
    bullet("All permission checks funnel through a single PermissionService with a 5-minute in-memory cache — cache invalidation must be triggered explicitly on every permission-affecting write")
    bullet("Microsoft Entra ID integration is fully optional and self-disabling: the OpenID Connect scheme is only registered when both AzureAd:ClientId and AzureAd:TenantId are configured, since an unconfigured OIDC scheme would break authentication middleware for the entire application")

    h2("2.6 Assumptions and Dependencies")
    bullet("The organization's shift structure fits the fixed 5-group (A/B/C/D/F) x 3-shift (Morning/Evening/Night) + Off rotation model")
    bullet("Reliable outbound connectivity is available for optional integrations (Entra ID, SMTP, Twilio WhatsApp, OpenAI/Azure OpenAI, Azure Speech) — the application degrades gracefully (features simply unavailable) when these are not configured, except that a misconfigured (partially-configured) Entra block is treated as “not configured” and fully disabled")
    bullet("Server and client clocks are reasonably synchronized; shift time windows are computed against a hard-coded UTC+3 offset")

    page_break()
