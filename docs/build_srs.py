# -*- coding: utf-8 -*-
"""Builds SRS.docx for ShiftFlow (solution3). Run with: python build_srs.py"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(HERE, "srs_assets")
OUT = os.path.join(HERE, "SRS.docx")

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT = RGBColor(0x0D, 0x6E, 0xFD)
GRAY = RGBColor(0x5A, 0x5A, 0x5A)

doc = Document()

# ---------- base style tuning ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)

for i, size in zip((1, 2, 3, 4), (22, 16, 13, 11.5)):
    st = doc.styles[f"Heading {i}"]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = NAVY
    st.font.bold = True
    st.paragraph_format.space_before = Pt(18 if i == 1 else 12)
    st.paragraph_format.space_after = Pt(8 if i == 1 else 6)


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)


def add_toc(document):
    p = document.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t"); fld_text.text = "Right-click and choose 'Update Field' to build the Table of Contents."
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_sep)
    r2 = p.add_run(); r2._r.append(fld_text)
    run2 = p.add_run(); run2._r.append(fld_end)


def h1(text, numbered=True):
    p = doc.add_heading(text, level=1)
    return p


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def h4(text):
    return doc.add_heading(text, level=4)


def para(text, italic=False, bold=False, size=None, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet" if level == 0 else "List Bullet 2")
    return p


def numbered(text):
    return doc.add_paragraph(text, style="List Number")


def add_diagram(filename, caption, width_in=6.3):
    path = os.path.join(ASSETS, filename)
    if not os.path.exists(path):
        para(f"[Missing diagram: {filename}]", italic=True, color=RGBColor(0xC0, 0x00, 0x00))
        return
    doc.add_picture(path, width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = GRAY


def simple_table(headers, rows, widths=None, header_color="1B2A4A"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(htext)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        set_cell_shading(hdr_cells[i], header_color)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def page_break():
    doc.add_page_break()


def section_break_with_footer():
    """Adds a new section (for page numbering after the title page/TOC)."""
    doc.add_section()


# ---------- footer with page numbers ----------
section = doc.sections[0]
section.footer.is_linked_to_previous = False
footer_p = section.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.add_run("ShiftFlow SRS — Page ")
add_page_number_field(footer_p)

# ---------- assemble content ----------
ctx = {
    "doc": doc, "h1": h1, "h2": h2, "h3": h3, "h4": h4,
    "para": para, "bullet": bullet, "numbered": numbered,
    "add_diagram": add_diagram, "simple_table": simple_table,
    "page_break": page_break, "add_toc": add_toc,
    "NAVY": NAVY, "ACCENT": ACCENT, "GRAY": GRAY,
}

import importlib.util


def load_module(name, path):
    spec = importlib.util.spec_from_file_location(name, path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


part1 = load_module("srs_content_1", os.path.join(HERE, "srs_content_1.py"))
part2 = load_module("srs_content_2", os.path.join(HERE, "srs_content_2.py"))
part3 = load_module("srs_content_3", os.path.join(HERE, "srs_content_3.py"))
part4 = load_module("srs_content_4", os.path.join(HERE, "srs_content_4.py"))
part5 = load_module("srs_content_5", os.path.join(HERE, "srs_content_5.py"))

part1.build_part1(ctx)
part2.build_part2(ctx)
part3.build_part3(ctx)
part4.build_part4(ctx)
part5.build_part5(ctx)

doc.save(OUT)
print(f"Saved SRS to {OUT}")
